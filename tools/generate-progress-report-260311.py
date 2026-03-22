"""Generate Vietnamese progress report (11-03-2026) in .docx format."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p


def make_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    return table


# ══════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BÁO CÁO DỰ ÁN')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('"XÁC ĐỊNH GIÁ TRỊ BOD VÀ ĐỘ ĐỘC CHO CẢM BIẾN SINH HỌC"')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

add_para('Đơn vị: Sinh Học.')
add_para('Thời gian thực hiện: Từ 16/11/2025 đến 11/03/2026.')
add_para('Thiết bị triển khai: Thiết bị đo sinh học.')
add_para('Loại dữ liệu nhận diện: GGA, GGA-metal.')

# ══════════════════════════════════════════════════
# I. MỤC TIÊU DỰ ÁN
# ══════════════════════════════════════════════════
add_heading_styled('I. MỤC TIÊU DỰ ÁN', level=1)
add_para(
    'Mục tiêu chính của dự án này là phát triển một mô hình học máy có khả năng '
    'nhận diện nguyên nhân gây độc và nồng độ độc trong các mẫu sinh học.'
)
add_para(
    'Mô hình nhằm hỗ trợ tự động hóa và cải thiện độ chính xác trong phân tích nồng độ, '
    'điều này rất quan trọng trong giám sát môi trường và các ứng dụng công nghiệp.'
)

# ══════════════════════════════════════════════════
# II. CÁC THIẾT BỊ VÀ DỮ LIỆU
# ══════════════════════════════════════════════════
add_heading_styled('II. CÁC THIẾT BỊ VÀ DỮ LIỆU', level=1)

add_para('Thiết bị', bold=True)
add_para('Thiết bị đo sinh học để thu thập dữ liệu tín hiệu.')
add_para(
    'Môi trường tính toán: Python 3.9 với TensorFlow, Scikit-learn, CatBoost, SciPy '
    'và các thư viện khác phục vụ xử lý dữ liệu và huấn luyện mô hình.'
)

add_para('Dữ liệu', bold=True)
add_para('Nguồn:', bold=True)
add_para(
    'Dữ liệu sinh học của các mẫu phân thành 2 dạng biến đổi nồng độ do môi trường '
    'và biến đổi nồng độ do vi sinh vật thu nhập từ các tệp .txt trong 2 thư mục "GGA" và "GGA-metal"'
)
add_para(
    'Dữ liệu các giá trị hpeak và các điểm đặc biệt như DO_min, DO_in thu nhập từ '
    'các tệp .xlsx trong thư mục "File excel" (dùng làm ground truth để đánh giá thuật toán)'
)

add_para('Cấu trúc:', bold=True)
add_para('Mỗi tệp .txt chứa: Thời gian đo (s) và Giá trị BOD (V).', indent=True)
add_para(
    'Metadata: Dữ liệu được xử lý và xuất thành các tệp CSV '
    '(vd: metadata-gga-txt.csv) để dễ dàng sử dụng.',
    indent=True,
)

add_para('Bộ dữ liệu đánh giá:', bold=True)
add_para('Test: 14 mẫu, 681 peak ground truth', indent=True)
add_para('GGA: 134 mẫu, 1.793 peak ground truth', indent=True)
add_para('Metal: 348 mẫu, 5.254 peak ground truth', indent=True)
add_para('HH: 89 mẫu, 1.252 peak ground truth', indent=True)

# ══════════════════════════════════════════════════
# III. CÔNG VIỆC ĐÃ THỰC HIỆN
# ══════════════════════════════════════════════════
add_heading_styled('III. CÔNG VIỆC ĐÃ THỰC HIỆN', level=1)

add_para(
    'Trong giai đoạn này, nhóm đã phát triển thành công thuật toán trích xuất đặc trưng '
    'hoàn toàn tự động từ file .txt gốc, thay thế hoàn toàn việc phụ thuộc vào file Excel '
    'từ chuyên gia. Đồng thời tích hợp thuật toán vào pipeline production.'
)

# 1. Algorithm
add_para('1. Phát triển thuật toán Adaptive Peak Extraction', bold=True)

add_para('1.1 Kiến trúc thuật toán Two-Pass HH Detection', bold=True)
add_para('Thuật toán hoạt động theo 2 bước:', indent=True)
add_para(
    'Bước 1 (Pass 1): Trích xuất với tham số non-HH (conservative) -> tính thống kê DDO',
    indent=True,
)
add_para(
    'Bước 2 (Pass 2): Nếu DDO statistics cho thấy tín hiệu thuộc loại HH '
    '(DDO P90 > 12mV, DDO max > 20mV), thuật toán tự động re-extract với tham số HH riêng biệt',
    indent=True,
)

add_para('1.2 Các thành phần chính của thuật toán', bold=True)
add_para(
    'Cycle Detection: Autocorrelation + FFT fallback để tự động xác định chu kỳ tín hiệu',
    indent=True,
)
add_para(
    'Signal Classification: Phân loại tín hiệu thành low/medium/high variance',
    indent=True,
)
add_para(
    'Minima Detection: Local prominence filtering + HH pre-filter (bottom 25% of robust range)',
    indent=True,
)
add_para(
    'Plateau (DOin) Detection: Gradient-based drop detection -> stable value collection '
    '-> IQR filtering -> bias correction',
    indent=True,
)
add_para(
    'Noise Filtering: Adaptive DDO threshold + bimodal gap detection cho HH signals',
    indent=True,
)

add_para('1.3 Tham số đã tối ưu (sau 485+ cấu hình thử nghiệm)', bold=True)
make_table(
    ['Tham số', 'Giá trị', 'Mô tả'],
    [
        ['safety', '5', 'Khoảng cách an toàn từ drop start'],
        ['hh_smooth', '21', 'Smoothing cho DOin estimation (HH signals)'],
        ['non_hh_smooth', '19', 'Smoothing cho DOin estimation (non-HH signals)'],
        ['hh_lookback', '0.50', 'Tỷ lệ lookback cho HH (50% cycle)'],
        ['non_hh_lookback', '0.75', 'Tỷ lệ lookback cho non-HH (75% cycle)'],
        ['stab_mult', '0.012', 'Gradient stability threshold multiplier'],
        ['hh_bias', '+0.04 mV', 'Bias correction cho HH underestimation'],
        ['non_hh_bias', '+0.05 mV', 'Bias correction cho non-HH underestimation'],
        ['hh_grad_smooth', '7', 'Gradient smoothing (HH)'],
        ['non_hh_grad_smooth', '11', 'Gradient smoothing (non-HH)'],
        ['pos_threshold', '100', 'Position constraint cho peak matching'],
    ],
)

# 2. Bug fixes
doc.add_paragraph()
add_para('2. Phân tích lỗi và sửa bug đánh giá', bold=True)

add_para('2.1 Phát hiện và sửa 3 bug nghiêm trọng trong hệ thống đánh giá:', bold=True)
add_para(
    'Double-matching bug: Peak ground truth bị match trùng với cùng một extracted peak '
    '-> sửa bằng used_ext_indices set',
    indent=True,
)
add_para(
    'Key collision bug: Tên mẫu bị trùng key khi normalize '
    '-> sửa bằng individual metal regex, BOD values, 7-digit dates',
    indent=True,
)
add_para(
    'Cross-cycle false match: Peak match nhầm sang cycle khác do thiếu position constraint '
    '-> thêm position constraint 100 điểm',
    indent=True,
)

add_para('2.2 Phân tích DOin Error Distribution', bold=True)
add_para(
    '78% DOin failures trên tất cả non-HH datasets đều bị under-estimated '
    '(ước lượng thấp hơn ground truth)',
    indent=True,
)
add_para('Mean signed error: -0.451mV -> systematic directional bias', indent=True)
add_para(
    'Giải pháp: Áp dụng bias correction +0.05mV cho non-HH, +0.04mV cho HH '
    '-> cải thiện Metal từ 77.2% lên 80.1% (tại 0.3mV tolerance)',
    indent=True,
)

# 3. ML refinement
doc.add_paragraph()
add_para('3. Thử nghiệm ML Refinement (đã bác bỏ)', bold=True)
add_para('Thử nghiệm 2 mô hình ML để tinh chỉnh DOin estimation:', indent=True)
add_para(
    'Gradient Boosting Regressor (300 trees): -1.1% so với uniform +0.05 bias',
    indent=True,
)
add_para('Random Forest (200 trees): -0.7% so với uniform +0.05 bias', indent=True)
add_para(
    'Kết luận: ML introduces prediction noise tệ hơn constant correction. '
    'Uniform bias +0.05mV là tối ưu.',
    indent=True,
)

# 4. Pipeline integration
doc.add_paragraph()
add_para('4. Tích hợp Pipeline End-to-End', bold=True)
add_para('Hoàn thành tích hợp thuật toán vào pipeline production:')
add_para(
    'Tạo module src/peak_extractor.py (449 dòng) — core algorithm production-ready',
    indent=True,
)
add_para(
    'Thêm hàm extract_peaks_from_txt() trong src/utils.py — bridge TXT -> peak extraction -> DataFrame',
    indent=True,
)
add_para(
    'Cập nhật app.py — loại bỏ hoàn toàn phụ thuộc Excel, chuyển sang TXT-only flow 5 bước:',
    indent=True,
)
add_para('Bước 1: Upload file .txt', indent=True)
add_para('Bước 2: Extract Peaks (thuật toán adaptive)', indent=True)
add_para('Bước 3: LSTM Prediction (dự đoán DO)', indent=True)
add_para('Bước 4: Classification (CatBoost — GGA/GGA-metal)', indent=True)
add_para('Bước 5: Toxicity Calculation (tính nồng độ độc)', indent=True)

doc.add_paragraph()
add_para('Pipeline mới:', bold=True)
add_para(
    'Upload TXT (UTF-16) -> extract_peaks_from_txt() -> peaks DataFrame -> Classification / Toxicity'
)
add_para(
    '                                    -> process_and_predict_lstm() (LSTM prediction độc lập)'
)

# ══════════════════════════════════════════════════
# IV. KẾT QUẢ ĐẠT ĐƯỢC
# ══════════════════════════════════════════════════
add_heading_styled('IV. KẾT QUẢ ĐẠT ĐƯỢC', level=1)

add_para('Bảng 1 - Kết quả trích xuất đặc trưng (Adaptive Peak Extraction)', bold=True)
make_table(
    ['Dataset', 'Số mẫu', 'Số peak GT', 'Accuracy (0.2mV)', 'Accuracy (0.3mV)'],
    [
        ['Test', '14', '681', '88.1%', '93.0%'],
        ['GGA', '134', '1.793', '79.9%', '85.3%'],
        ['Metal', '348', '5.254', '73.6%', '80.1%'],
        ['HH', '89', '1.252', '73.7%', '84.1%'],
    ],
)

doc.add_paragraph()
add_para('Bảng 2 - So sánh với kết quả báo cáo trước (18/11/2025)', bold=True)
make_table(
    ['Phương pháp', 'Sai số DDO', 'Tỷ lệ sai >0.2mV', 'Ghi chú'],
    [
        [
            'Rule-based cũ (11/2025)',
            '0.1 - 3.5 mV',
            '>90%',
            'Fixed interval, không adaptive',
        ],
        ['XGBoost cũ (11/2025)', '2.0 - 5.0 mV', '100%', 'ML predict start/interval'],
        [
            'Adaptive Two-Pass (03/2026)',
            '<0.3mV (80-93%)',
            '7-27%',
            '485+ configs, bias correction',
        ],
    ],
)

doc.add_paragraph()
add_para(
    'Cải thiện đáng kể: Từ >90% tỷ lệ sai ở báo cáo trước -> chỉ còn 7-27% tỷ lệ sai '
    'tùy dataset. Đặc biệt, bộ dữ liệu Test đạt 93.0% accuracy tại 0.3mV tolerance.'
)

doc.add_paragraph()
add_para('Bảng 3 - Kết quả mô hình phân loại CatBoost', bold=True)
make_table(
    ['Models', 'Parameters', 'Accuracy'],
    [
        [
            'CatBoost',
            'iterations=1000, learning_rate=0.1, depth=6, l2_leaf_reg=3, '
            'auto_class_weights="Balanced", eval_metric="F1", od_type="Iter", od_wait=50',
            '0.81',
        ],
    ],
)

doc.add_paragraph()
add_para('Bảng 4 - Kết quả mô hình LSTM', bold=True)
make_table(
    ['Models', 'Parameters', 'Train RMSE', 'Test RMSE'],
    [
        [
            'Encoder-Decoder LSTM',
            'lookback=7, LSTM units=[128,128,64], activation=relu, '
            'optimizer=Adam, loss=Huber, EarlyStopping(patience=10)',
            '0.2121',
            '1.4268',
        ],
    ],
)

# ══════════════════════════════════════════════════
# V. KHÓ KHĂN
# ══════════════════════════════════════════════════
add_heading_styled('V. KHÓ KHĂN', level=1)
add_para(
    'Trong giai đoạn tối ưu hóa thuật toán trích xuất, nhóm gặp các khó khăn sau:'
)

add_para('Physical noise ceiling:', bold=True)
add_para(
    'Tín hiệu HH có nhiễu plateau ~0.17mV, gần với tolerance 0.2mV -> giới hạn vật lý '
    'khoảng 74% accuracy cho HH tại 0.2mV. Đây là giới hạn của tín hiệu chứ không phải thuật toán.',
    indent=True,
)

add_para('Metal structural limit:', bold=True)
add_para(
    'Tín hiệu Metal có hình dạng plateau không đều (irregular plateau shapes), '
    'khiến việc xác định DOin chính xác trở nên khó khăn. '
    'Accuracy Metal đạt ceiling ~80.1% tại 0.3mV.',
    indent=True,
)

add_para('Systematic DOin under-estimation:', bold=True)
add_para(
    '78% DOin failures đều bị ước lượng thấp. Sau khi áp dụng bias correction +0.05mV, '
    'accuracy cải thiện nhưng không thể hoàn toàn loại bỏ do tính chất vật lý của tín hiệu.',
    indent=True,
)

add_para('ML refinement không hiệu quả:', bold=True)
add_para(
    'Cả 2 mô hình ML (GBR, RF) đều cho kết quả kém hơn uniform bias. '
    'ML introduces prediction noise tệ hơn constant correction trên 8.384 training samples.',
    indent=True,
)

add_para('Dữ liệu mất cân bằng:', bold=True)
add_para(
    'Mẫu GGA Metal chiếm tỷ lệ cao (3/4), ảnh hưởng đến '
    'khả năng tổng quát hóa của mô hình phân loại.',
    indent=True,
)

# ══════════════════════════════════════════════════
# VI. KIẾN NGHỊ
# ══════════════════════════════════════════════════
add_heading_styled('VI. KIẾN NGHỊ', level=1)
add_para(
    'Đề xuất được cung cấp thêm dữ liệu về mẫu GGA, nhằm cân bằng lại phân bố '
    'dữ liệu huấn luyện và cải thiện hiệu suất, độ chính xác của mô hình phân loại.'
)
add_para(
    'Đề xuất chuyên gia cung cấp smoothed curve + baseline từ OriginLab để có thể '
    'so sánh trực tiếp cách tính DOin giữa thuật toán và chuyên gia, '
    'từ đó có thể cải thiện thêm accuracy nếu cần.'
)
add_para(
    'Xem xét chấp nhận tolerance 0.3mV thay vì 0.2mV cho production use, '
    'vì accuracy cải thiện đáng kể (từ 73-88% lên 80-93%).'
)

# ══════════════════════════════════════════════════
# BẢNG TIẾN ĐỘ
# ══════════════════════════════════════════════════
add_heading_styled('BẢNG TIẾN ĐỘ DỰ ÁN', level=1)
make_table(
    ['Thời gian', 'Công việc', 'Nội dung', 'Ghi chú'],
    [
        [
            '10/2024',
            'Tiếp nhận và xử lý dữ liệu',
            '2 bộ dữ liệu GGA và GGA-metal, hơn 1 triệu mẫu. '
            'Chuỗi thời gian, không đồng nhất số lượng điểm dữ liệu.',
            '',
        ],
        [
            '11/2024 - 12/2024',
            'Xây dựng mô hình ML',
            'Các mô hình học kém trên bộ dữ liệu GGA',
            'Không thành công (mô hình không có yếu tố thời gian)',
        ],
        [
            '01/2025 - 03/2025',
            'Chuyển sang Deep Learning LSTM',
            'Mô hình cải thiện, predict tốt trên GGA',
            'Đạt kết quả theo yêu cầu',
        ],
        [
            '03/2025 - 07/2025',
            'Xây dựng pipeline end-to-end',
            'Trích xuất đặc trưng + pipeline phân loại GGA/GGA-metal',
            'Phân loại cần cải tiến',
        ],
        [
            '07/2025 - 10/2025',
            'Fallback trích xuất đặc trưng',
            'API OriginPro Lab',
            'Không thành công',
        ],
        [
            '11/2025',
            'Rule-based + ML extraction',
            'Rule-based: DDO error 0.1-3.5mV (>90% fail). '
            'XGBoost: 2-5mV (100% fail)',
            'Chưa đạt yêu cầu',
        ],
        [
            '11/2025 - 02/2026',
            'Adaptive Peak Extraction',
            'Two-pass HH detection, 485+ configs, sửa 3 bug đánh giá, '
            'bias correction +0.05mV. Test 93.0%, GGA 85.3%, Metal 80.1%, HH 84.1% @0.3mV',
            'Đạt kết quả đột phá',
        ],
        [
            '02/2026 - 03/2026',
            'Tích hợp pipeline production',
            'Tạo src/peak_extractor.py, cập nhật app.py thành TXT-only 5-step pipeline, '
            'loại bỏ phụ thuộc Excel hoàn toàn',
            'Pipeline hoàn chỉnh',
        ],
    ],
)

# Save
output_path = 'plans/reports/BÁO CÁO SINH HỌC (11-03-2026).docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
